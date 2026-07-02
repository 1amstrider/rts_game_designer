"""
Convert all .md files in design/ folder to .pdf and .docx
Usage: python design/convert_docs.py
"""

import os
import re
from pathlib import Path
from markdown import markdown

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("Warning: python-docx not available, skipping .docx generation")

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
    print("Warning: reportlab not available, skipping .pdf generation")


def parse_markdown_lines(md_text):
    """Parse markdown into structured lines."""
    lines = []
    current_table = []
    in_table = False
    
    for line in md_text.split('\n'):
        stripped = line.strip()
        
        # Table handling
        if stripped.startswith('|') and stripped.endswith('|'):
            if not in_table:
                in_table = True
                current_table = []
            # Skip separator rows (---)
            if not re.match(r'\|[-:\s|]+\|', stripped):
                cells = [c.strip() for c in stripped[1:-1].split('|')]
                current_table.append(cells)
            continue
        elif in_table:
            in_table = False
            if current_table:
                lines.append(('table', current_table))
            current_table = []
        
        # Empty line
        if not stripped:
            continue
        
        # Headers
        if stripped.startswith('# '):
            lines.append(('h1', stripped[2:]))
        elif stripped.startswith('## '):
            lines.append(('h2', stripped[3:]))
        elif stripped.startswith('### '):
            lines.append(('h3', stripped[4:]))
        elif stripped.startswith('#### '):
            lines.append(('h4', stripped[5:]))
        # Horizontal rule
        elif stripped == '---' or stripped == '***':
            lines.append(('hr', ''))
        # Lists
        elif stripped.startswith('- ') or stripped.startswith('* '):
            lines.append(('ul', stripped[2:]))
        elif re.match(r'^\d+\.\s', stripped):
            lines.append(('ol', re.sub(r'^\d+\.\s', '', stripped)))
        # Blockquote
        elif stripped.startswith('> '):
            lines.append(('quote', stripped[2:]))
        # Regular paragraph
        else:
            # Handle inline formatting: bold, italic, code
            lines.append(('p', stripped))
    
    # Handle table at end of file
    if in_table and current_table:
        lines.append(('table', current_table))
    
    return lines


def md_to_docx(md_path, docx_path):
    """Convert markdown file to Word document."""
    if not HAS_DOCX:
        return False
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    md_text = md_path.read_text(encoding='utf-8')
    lines = parse_markdown_lines(md_text)
    
    for line_type, content in lines:
        if line_type == 'h1':
            p = doc.add_heading(content, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line_type == 'h2':
            doc.add_heading(content, level=2)
        elif line_type == 'h3':
            doc.add_heading(content, level=3)
        elif line_type == 'h4':
            doc.add_heading(content, level=4)
        elif line_type == 'hr':
            doc.add_paragraph('_' * 50)
        elif line_type == 'ul':
            p = doc.add_paragraph(content, style='List Bullet')
        elif line_type == 'ol':
            p = doc.add_paragraph(content, style='List Number')
        elif line_type == 'quote':
            p = doc.add_paragraph(content)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.25)
            run = p.runs[0] if p.runs else p.add_run()
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        elif line_type == 'table':
            if not content:
                continue
            table = doc.add_table(rows=1, cols=len(content[0]))
            table.style = 'Table Grid'
            
            # Header row
            for i, cell_text in enumerate(content[0]):
                cell = table.rows[0].cells[i]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Data rows
            for row_data in content[1:]:
                row = table.add_row()
                for i, cell_text in enumerate(row_data):
                    row.cells[i].text = cell_text
            
            doc.add_paragraph()  # Spacing after table
        elif line_type == 'p':
            # Handle bold and italic
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                else:
                    p.add_run(part)
    
    doc.save(docx_path)
    return True


def md_to_pdf(md_path, pdf_path):
    """Convert markdown file to PDF."""
    if not HAS_PDF:
        return False
    
    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=18
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    styles.add(ParagraphStyle(
        name='CustomH1',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=TA_CENTER,
        textColor='#2c3e50'
    ))
    styles.add(ParagraphStyle(
        name='CustomH2',
        parent=styles['Heading2'],
        fontSize=18,
        spaceAfter=12,
        spaceBefore=12,
        textColor='#34495e'
    ))
    styles.add(ParagraphStyle(
        name='CustomH3',
        parent=styles['Heading3'],
        fontSize=14,
        spaceAfter=10,
        spaceBefore=10,
        textColor='#7f8c8d'
    ))
    styles.add(ParagraphStyle(
        name='CustomQuote',
        parent=styles['Normal'],
        fontSize=10,
        leftIndent=20,
        rightIndent=20,
        textColor='#555555',
        fontName='Helvetica-Oblique'
    ))
    styles.add(ParagraphStyle(
        name='CustomCode',
        parent=styles['Normal'],
        fontName='Courier',
        fontSize=9,
        textColor='#c0392b'
    ))
    
    story = []
    md_text = md_path.read_text(encoding='utf-8')
    lines = parse_markdown_lines(md_text)
    
    for line_type, content in lines:
        if line_type == 'h1':
            story.append(Paragraph(content, styles['CustomH1']))
            story.append(Spacer(1, 0.2*inch))
        elif line_type == 'h2':
            story.append(Paragraph(content, styles['CustomH2']))
            story.append(Spacer(1, 0.1*inch))
        elif line_type == 'h3':
            story.append(Paragraph(content, styles['CustomH3']))
            story.append(Spacer(1, 0.05*inch))
        elif line_type == 'hr':
            story.append(Spacer(1, 0.2*inch))
        elif line_type == 'ul':
            story.append(Paragraph(f'• {content}', styles['Normal']))
        elif line_type == 'ol':
            story.append(Paragraph(f'{content}', styles['Normal']))
        elif line_type == 'quote':
            story.append(Paragraph(content, styles['CustomQuote']))
        elif line_type == 'table':
            # Create table data
            if content:
                from reportlab.platypus import Table, TableStyle
                from reportlab.lib import colors
                
                t = Table(content, colWidths=[2*inch]*len(content[0]))
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#34495e')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#ecf0f1')),
                    ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#bdc3c7')),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#ecf0f1')]),
                ]))
                story.append(t)
                story.append(Spacer(1, 0.1*inch))
        elif line_type == 'p':
            # Convert markdown to HTML for reportlab
            html_content = markdown(content, extensions=['tables'])
            # Strip HTML tags for basic rendering
            text = re.sub(r'<[^>]+>', '', html_content)
            story.append(Paragraph(text, styles['Normal']))
    
    doc.build(story)
    return True


def main():
    design_dir = Path(__file__).parent
    
    md_files = list(design_dir.glob('*.md'))
    
    if not md_files:
        print("No .md files found in design/ folder")
        return
    
    print(f"Found {len(md_files)} markdown files")
    
    for md_file in md_files:
        print(f"\nConverting: {md_file.name}")
        
        # Convert to DOCX
        if HAS_DOCX:
            docx_path = md_file.with_suffix('.docx')
            if md_to_docx(md_file, docx_path):
                print(f"  [OK] Created: {docx_path.name}")
            else:
                print(f"  [FAIL] Failed to create DOCX")
        
        # Convert to PDF
        if HAS_PDF:
            pdf_path = md_file.with_suffix('.pdf')
            if md_to_pdf(md_file, pdf_path):
                print(f"  [OK] Created: {pdf_path.name}")
            else:
                print(f"  [FAIL] Failed to create PDF")
    
    print("\nDone! Check the design/ folder for generated files.")


if __name__ == '__main__':
    main()
